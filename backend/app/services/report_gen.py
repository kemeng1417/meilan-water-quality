"""Word 报告生成 —— 横向排版，标准限值单独参考行"""
import io
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

import os
from app.models import TestRecord, TestDetail, SamplePoint, Indicator, WaterType, StandardLimit, Photo
from app.config import DATA_DIR

PHOTOS_DIR = os.path.join(DATA_DIR, "photos")


def generate_word_report(db: Session, record_id: int) -> io.BytesIO:
    record = db.query(TestRecord).filter(TestRecord.id == record_id).first()
    water_type = db.query(WaterType).filter(WaterType.id == record.water_type_id).first()
    details = db.query(TestDetail).filter(TestDetail.record_id == record_id).all()
    indicators = db.query(Indicator).order_by(Indicator.display_order).all()
    sample_points = db.query(SamplePoint).filter(
        SamplePoint.id.in_(list(set(d.sample_point_id for d in details)))
    ).order_by(SamplePoint.area, SamplePoint.sort_order).all()

    # For combined reports, load both 出厂水 and 末梢水 limits
    if record.water_type_id == 4:
        limits_1 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == 1).all()}
        limits_2 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == 2).all()}
        limits = limits_1
    else:
        limits_1 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == record.water_type_id).all()}
        limits_2 = None
        limits = limits_1

    matrix = {}
    for d in details:
        matrix.setdefault(d.sample_point_id, {})[d.indicator_id] = d

    doc = Document()

    # ── 横向页面 ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.0)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 标题 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run('海口美兰机场供水站水质工作反馈')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 信息行 ──
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(8)
    info_text = f"化验日期：{record.test_date}    报告日期：{record.report_date}    执行标准：{water_type.standard_code if water_type else '—'}"
    ir = info.add_run(info_text)
    ir.font.size = Pt(12)
    ir.font.name = '宋体'
    ir.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 表格（序号+采样点+N指标+结论） ──
    cols = 3 + len(indicators)
    table = doc.add_table(rows=2 + len(sample_points), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 列宽设置
    widths = [Cm(0.7), Cm(3.5)] + [Cm(2.0)] * len(indicators) + [Cm(1.5)]

    def _set_cell(cell, text, font_size=12, bold=False, bg=None, color=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if color:
            run.font.color.rgb = color
        if bg:
            from docx.oxml import OxmlElement
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg)
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

    # ── 表头行 ──
    hdr_bg = 'D9E1F2'
    _set_cell(table.rows[0].cells[0], '序号', 12, True, hdr_bg)
    _set_cell(table.rows[0].cells[1], '采样点', 12, True, hdr_bg)
    for i, ind in enumerate(indicators):
        txt = f"{ind.name}\n({ind.unit})" if ind.unit else ind.name
        _set_cell(table.rows[0].cells[2 + i], txt, 10, True, hdr_bg)
    _set_cell(table.rows[0].cells[2 + len(indicators)], '结论', 12, True, hdr_bg)

    # ── 标准参考行 ──
    _set_cell(table.rows[1].cells[0], '', 10, False, 'F0F5FA')
    _set_cell(table.rows[1].cells[1], '标准限值', 10, True, 'F0F5FA')
    for i, ind in enumerate(indicators):
        lim = limits_1.get(ind.id)
        lim2 = limits_2.get(ind.id) if limits_2 else None
        txt = _fmt_limit_short(ind, lim, lim2)
        _set_cell(table.rows[1].cells[2 + i], txt, 9, False, 'F0F5FA')
    _set_cell(table.rows[1].cells[2 + len(indicators)], '', 10, False, 'F0F5FA')

    # ── 数据行 ──
    for r, sp in enumerate(sample_points):
        row = table.rows[2 + r]
        _set_cell(row.cells[0], str(r + 1), 12)
        _set_cell(row.cells[1], sp.name, 12)

        pt_details = matrix.get(sp.id, {})
        has_abnormal = any(d.is_abnormal for d in pt_details.values())
        all_filled = bool(pt_details) and all(d.value_text for d in pt_details.values())

        for i, ind in enumerate(indicators):
            cell = row.cells[2 + i]
            detail = pt_details.get(ind.id)
            if detail and detail.value_text:
                is_fail = detail.is_qualified is False
                bg = 'FCE4D6' if is_fail else ('E2EFDA' if detail.is_qualified else None)
                color = RGBColor(0xFF, 0x00, 0x00) if is_fail else None
                _set_cell(cell, detail.value_text, 12, is_fail, bg, color)
            else:
                _set_cell(cell, '—', 12, False, None, None)

        # 结论列
        conc_col_idx = 2 + len(indicators)
        if has_abnormal:
            _set_cell(row.cells[conc_col_idx], '不合格', 12, True, 'FCE4D6', RGBColor(0xFF, 0x00, 0x00))
        elif all_filled:
            _set_cell(row.cells[conc_col_idx], '合格', 12, False, 'E2EFDA', RGBColor(0x16, 0xA3, 0x4A))
        else:
            _set_cell(row.cells[conc_col_idx], '—', 12, False, None, None)

    # ── 异常汇总 ──
    abnormal_items = [d for d in details if d.is_abnormal]
    doc.add_paragraph()
    conc = doc.add_paragraph()
    conc.paragraph_format.space_before = Pt(2)
    if record.conclusion:
        cr = conc.add_run(record.conclusion)
        if record.is_abnormal:
            cr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    elif abnormal_items:
        pts = list({d.sample_point_id for d in abnormal_items})
        pt_names = [sp.name for sp in sample_points if sp.id in pts]
        cr = conc.add_run(f"超标项: {len(abnormal_items)} 项 | 超标点位: {'、'.join(pt_names)}")
        cr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    else:
        cr = conc.add_run(f"结论：本次检测项目全部合格，符合{water_type.standard_code if water_type else '相关'}标准要求。")
    cr.font.size = Pt(12)
    cr.font.name = '宋体'
    cr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 签名 ──
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(16)
    sr = sig.add_run(f"化验员：{record.tester}              审核人：{record.reviewer or '___________'}              日期：{record.report_date}")
    sr.font.size = Pt(12)
    sr.font.name = '宋体'
    sr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 照片 ──
    photos = db.query(Photo).filter(Photo.record_id == record_id).all()
    if photos:
        doc.add_paragraph()
        photo_title = doc.add_paragraph()
        photo_title.paragraph_format.space_before = Pt(10)
        pr = photo_title.add_run("现场照片：")
        pr.bold = True
        pr.font.size = Pt(12)
        pr.font.name = '宋体'
        pr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # Group photos by sample point
        photo_by_point: dict[int, list] = {}
        for p in photos:
            photo_by_point.setdefault(p.sample_point_id, []).append(p)

        sp_map = {sp.id: sp for sp in sample_points}
        for sp_id, pt_photos in photo_by_point.items():
            sp = sp_map.get(sp_id)
            sp_name = sp.name if sp else f"点位{sp_id}"
            lbl = doc.add_paragraph()
            lr = lbl.add_run(f"{sp_name}：")
            lr.font.size = Pt(11)
            lr.font.name = '宋体'
            lr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            for p in pt_photos:
                filepath = os.path.join(PHOTOS_DIR, p.filename)
                if os.path.exists(filepath):
                    try:
                        doc.add_picture(filepath, width=Cm(5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    except Exception:
                        doc.add_paragraph(f"[照片: {p.original_name}]")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fmt_limit_short(indicator, limit, limit2=None) -> str:
    """紧凑格式的限值文本，limit2 为联合报告时的第二组限值"""
    def _fmt_one(lim):
        if not lim:
            return None
        if lim.qual_check:
            return lim.qual_check
        if lim.min_value is not None and lim.max_value is not None:
            return f"{lim.min_value}-{lim.max_value}"
        if lim.max_value is not None:
            return f"≤{lim.max_value}"
        if lim.min_value is not None:
            return f"≥{lim.min_value}"
        return None

    t1 = _fmt_one(limit)
    if not limit2:
        return t1 or '—'
    t2 = _fmt_one(limit2)
    if not t2 or t1 == t2:
        return t1 or '—'
    return f"出厂:{t1} 末梢:{t2}"


def generate_html_report(db: Session, record_id: int) -> str:
    """生成自包含 HTML 报告"""
    record = db.query(TestRecord).filter(TestRecord.id == record_id).first()
    water_type = db.query(WaterType).filter(WaterType.id == record.water_type_id).first()
    details = db.query(TestDetail).filter(TestDetail.record_id == record_id).all()
    indicators = db.query(Indicator).order_by(Indicator.display_order).all()
    sample_points = db.query(SamplePoint).filter(
        SamplePoint.id.in_(list(set(d.sample_point_id for d in details)))
    ).order_by(SamplePoint.area, SamplePoint.sort_order).all()

    # For combined reports, load both 出厂水 and 末梢水 limits
    if record.water_type_id == 4:
        limits_1 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == 1).all()}
        limits_2 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == 2).all()}
        limits = limits_1
    else:
        limits_1 = {l.indicator_id: l for l in db.query(StandardLimit).filter(StandardLimit.water_type_id == record.water_type_id).all()}
        limits_2 = None
        limits = limits_1

    matrix = {}
    for d in details:
        matrix.setdefault(d.sample_point_id, {})[d.indicator_id] = d

    abnormal_items = [d for d in details if d.is_abnormal]

    # ── 生成指标表头 ──
    header_cells = ""
    for ind in indicators:
        txt = f"{ind.name}<br>({ind.unit})" if ind.unit else ind.name
        header_cells += f"<th>{txt}</th>"
    header_cells += "<th>结论</th>"

    # ── 标准参考行 ──
    limit_cells = ""
    for ind in indicators:
        lim = limits_1.get(ind.id)
        lim2 = limits_2.get(ind.id) if limits_2 else None
        txt = _fmt_limit_short(ind, lim, lim2)
        limit_cells += f"<td class='limit'>{txt}</td>"
    limit_cells += "<td class='limit'></td>"

    # ── 数据行 ──
    data_rows = ""
    for r, sp in enumerate(sample_points):
        cells = f"<td>{r + 1}</td><td class='point-name'>{sp.name}</td>"
        pt_details = matrix.get(sp.id, {})
        has_abnormal = any(d.is_abnormal for d in pt_details.values())
        all_filled = bool(pt_details) and all(d.value_text for d in pt_details.values())
        for ind in indicators:
            detail = pt_details.get(ind.id)
            if detail and detail.value_text:
                cls = 'fail' if detail.is_qualified is False else ('pass' if detail.is_qualified is True else '')
                cells += f"<td class='{cls}'>{detail.value_text}</td>"
            else:
                cells += f"<td class='empty'>—</td>"
        # 结论列
        if has_abnormal:
            cells += "<td class='fail'>不合格</td>"
        elif all_filled:
            cells += "<td class='pass'>合格</td>"
        else:
            cells += "<td class='empty'>—</td>"
        data_rows += f"<tr>{cells}</tr>"

    # ── 结论 ──
    if record.conclusion:
        conc_html = record.conclusion
        conc_cls = 'abnormal' if record.is_abnormal else ''
    elif abnormal_items:
        pts = list({d.sample_point_id for d in abnormal_items})
        pt_names = [sp.name for sp in sample_points if sp.id in pts]
        conc_html = f"超标项: {len(abnormal_items)} 项 | 超标点位: {'、'.join(pt_names)}"
        conc_cls = 'abnormal'
    else:
        conc_html = f"结论：本次检测项目全部合格，符合{water_type.standard_code if water_type else '相关'}标准要求。"
        conc_cls = ''

    # ── Photos (base64 embedded) ──
    photos_html = ""
    photos = db.query(Photo).filter(Photo.record_id == record_id).all()
    if photos:
        import base64
        photos_html += '<div class="photos"><h3>现场照片</h3>'
        photo_by_point: dict[int, list] = {}
        for p in photos:
            photo_by_point.setdefault(p.sample_point_id, []).append(p)

        sp_map = {sp.id: sp for sp in sample_points}
        for sp_id, pt_photos in photo_by_point.items():
            sp_name = sp_map[sp_id].name if sp_id in sp_map else f"点位{sp_id}"
            photos_html += f'<div class="photo-group"><p class="photo-label">{sp_name}：</p>'
            for p in pt_photos:
                filepath = os.path.join(PHOTOS_DIR, p.filename)
                if os.path.exists(filepath):
                    try:
                        with open(filepath, "rb") as f:
                            b64 = base64.b64encode(f.read()).decode()
                            ext = os.path.splitext(p.filename)[1].lower().replace('.', '')
                            mime = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"
                            photos_html += f'<img src="data:{mime};base64,{b64}" alt="{p.original_name}" class="photo" />'
                    except Exception:
                        photos_html += f'<span class="photo-missing">[照片: {p.original_name}]</span>'
            photos_html += '</div>'
        photos_html += '</div>'

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{record.record_no}</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: "SimSun", "宋体", "Microsoft YaHei", sans-serif; font-size: 12pt; color: #1e293b; padding: 30px 36px; }}
  h1 {{ text-align: center; font-size: 18pt; margin-bottom: 6px; font-weight: bold; }}
  .info {{ text-align: center; font-size: 11pt; margin-bottom: 18px; color: #475569; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 10pt; }}
  th, td {{ border: 1px solid #94a3b8; padding: 6px 5px; text-align: center; vertical-align: middle; }}
  th {{ background: #D9E1F2; font-weight: bold; font-size: 9.5pt; }}
  td.limit {{ background: #F0F5FA; font-size: 8.5pt; color: #475569; }}
  td.point-name {{ text-align: left; padding-left: 8px; }}
  td.pass {{ background: #E2EFDA; }}
  td.fail {{ background: #FCE4D6; color: #DC2626; font-weight: bold; }}
  td.empty {{ color: #94a3b8; }}
  .conclusion {{ margin-top: 16px; font-size: 11pt; line-height: 1.8; }}
  .conclusion.abnormal {{ color: #DC2626; font-weight: bold; }}
  .signature {{ margin-top: 28px; font-size: 11pt; }}
  .print-btn {{ position: fixed; top: 16px; right: 16px; padding: 10px 24px; background: #1e40af; color: #fff; border: none; border-radius: 6px; font-size: 13px; cursor: pointer; z-index: 1000; }}
  .print-btn:hover {{ background: #1e3a8a; }}
  .photos {{ margin-top: 20px; }}
  .photos h3 {{ font-size: 12pt; margin-bottom: 8px; }}
  .photo-group {{ margin-bottom: 10px; }}
  .photo-label {{ font-size: 10pt; font-weight: bold; margin-bottom: 4px; }}
  .photo {{ max-width: 200px; max-height: 150px; margin: 4px; border-radius: 6px; border: 1px solid #e8ecf1; }}
  .photo-missing {{ font-size: 9pt; color: #94a3b8; }}
  @media print {{
    body {{ padding: 10px 16px; }}
    @page {{ size: A4 landscape; margin: 1cm; }}
    .print-btn {{ display: none; }}
  }}
</style>
</head>
<body>
<button class="print-btn" onclick="window.print()">🖨 打印为PDF</button>
<h1>海口美兰机场供水站水质工作反馈</h1>
<p class="info">化验日期：{record.test_date} &emsp; 报告日期：{record.report_date} &emsp; 执行标准：{water_type.standard_code if water_type else '—'}</p>
<table>
<thead>
  <tr><th>序号</th><th>采样点</th>{header_cells}</tr>
</thead>
<tbody>
  <tr><td></td><td class="limit">标准限值</td>{limit_cells}</tr>
{data_rows}
</tbody>
</table>
<p class="conclusion {conc_cls}">{conc_html}</p>
<p class="signature">化验员：{record.tester} &emsp;&emsp;&emsp;&emsp; 审核人：{record.reviewer or '___________'} &emsp;&emsp;&emsp;&emsp; 日期：{record.report_date}</p>
{photos_html}
</body>
</html>"""
    return html
